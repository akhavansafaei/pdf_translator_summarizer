import os
from pathlib import Path
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class OutputWriter:
    """Handles writing processed content to various output formats."""

    def __init__(self, output_format: str = "docx", formatting_config: Optional[Dict[str, Any]] = None, output_language: str = "English"):
        """
        Initialize output writer.

        Args:
            output_format: Output format (docx, tex, txt, or md)
            formatting_config: Formatting configuration (font, text direction, etc.)
            output_language: Output language for auto-detecting RTL
        """
        self.output_format = output_format.lower()
        self.formatting_config = formatting_config or {}
        self.output_language = output_language

        # Determine text direction
        self.is_rtl = self._detect_rtl()

        # Get font settings
        font_config = self.formatting_config.get('font', {})
        self.font_name = font_config.get('name', 'B Nazanin')
        self.font_size = font_config.get('size', 13)
        self.heading_size = font_config.get('heading_size', 16)

    def _detect_rtl(self) -> bool:
        """
        Detect if text direction should be RTL based on configuration or language.

        Returns:
            True if RTL, False if LTR
        """
        text_direction = self.formatting_config.get('text_direction', 'auto')

        if text_direction == 'rtl':
            return True
        elif text_direction == 'ltr':
            return False
        else:  # auto
            # RTL languages
            rtl_languages = ['Persian', 'Arabic', 'Hebrew', 'Urdu', 'Farsi', 'Pashto', 'Kurdish']
            return any(lang.lower() in self.output_language.lower() for lang in rtl_languages)

    def write(self, content: str, output_path: Path, original_filename: str) -> bool:
        """
        Write content to file in specified format.

        Args:
            content: Processed content to write
            output_path: Directory path where file should be written
            original_filename: Original filename (without extension)

        Returns:
            True if successful, False otherwise
        """
        # Create output directory if it doesn't exist
        output_path.mkdir(parents=True, exist_ok=True)

        # Create output filename with appropriate extension
        output_file = output_path / f"{original_filename}.{self.output_format}"

        try:
            if self.output_format == "docx":
                return self._write_docx(content, output_file, original_filename)
            elif self.output_format == "tex":
                return self._write_tex(content, output_file)
            elif self.output_format == "txt":
                return self._write_txt(content, output_file)
            elif self.output_format == "md":
                return self._write_md(content, output_file)
            else:
                print(f"Error: Unsupported output format '{self.output_format}'")
                return False

        except Exception as e:
            print(f"Error writing file {output_file}: {str(e)}")
            return False

    def _set_rtl(self, paragraph):
        """
        Set RTL (right-to-left) text direction for a paragraph.

        Args:
            paragraph: The paragraph object to modify
        """
        pPr = paragraph._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    def _apply_font_and_rtl(self, paragraph, font_size=None):
        """
        Apply font and RTL settings to a paragraph.

        Args:
            paragraph: The paragraph to modify
            font_size: Font size in points (if None, uses self.font_size)
        """
        if font_size is None:
            font_size = self.font_size

        # Set alignment first
        if self.is_rtl:
            self._set_rtl(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Apply font to all runs in the paragraph
        for run in paragraph.runs:
            run.font.name = self.font_name
            run.font.size = Pt(font_size)
            # Set complex script font for RTL languages
            run._element.rPr.rFonts.set(qn('w:cs'), self.font_name)
            run._element.rPr.rFonts.set(qn('w:ascii'), self.font_name)
            run._element.rPr.rFonts.set(qn('w:hAnsi'), self.font_name)

    def _write_docx(self, content: str, output_file: Path, title: str) -> bool:
        """
        Write content to DOCX file with RTL support and Persian font.

        Args:
            content: Text content
            output_file: Path to output file
            title: Document title

        Returns:
            True if successful, False otherwise
        """
        try:
            doc = Document()

            # Add title - create empty heading first
            title_paragraph = doc.add_heading(level=1)
            title_run = title_paragraph.add_run(title)

            # Set title formatting
            title_run.font.name = self.font_name
            title_run.font.size = Pt(self.heading_size)
            title_run._element.rPr.rFonts.set(qn('w:cs'), self.font_name)
            title_run._element.rPr.rFonts.set(qn('w:ascii'), self.font_name)
            title_run._element.rPr.rFonts.set(qn('w:hAnsi'), self.font_name)

            if self.is_rtl:
                self._set_rtl(title_paragraph)
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add content - split into paragraphs
            paragraphs = content.split('\n\n')

            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it's a heading (starts with #)
                    if para_text.strip().startswith('#'):
                        # Remove markdown heading markers
                        heading_text = para_text.strip().lstrip('#').strip()
                        level = min(para_text.count('#'), 3)

                        # Create heading with text
                        heading = doc.add_heading(level=level + 1)
                        heading_run = heading.add_run(heading_text)

                        # Apply font
                        heading_run.font.name = self.font_name
                        heading_run.font.size = Pt(self.heading_size)
                        heading_run._element.rPr.rFonts.set(qn('w:cs'), self.font_name)
                        heading_run._element.rPr.rFonts.set(qn('w:ascii'), self.font_name)
                        heading_run._element.rPr.rFonts.set(qn('w:hAnsi'), self.font_name)

                        if self.is_rtl:
                            self._set_rtl(heading)
                            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        # Create paragraph with text
                        para = doc.add_paragraph()
                        para_run = para.add_run(para_text.strip())

                        # Apply font
                        para_run.font.name = self.font_name
                        para_run.font.size = Pt(self.font_size)
                        para_run._element.rPr.rFonts.set(qn('w:cs'), self.font_name)
                        para_run._element.rPr.rFonts.set(qn('w:ascii'), self.font_name)
                        para_run._element.rPr.rFonts.set(qn('w:hAnsi'), self.font_name)

                        # Set paragraph alignment and RTL
                        if self.is_rtl:
                            self._set_rtl(para)
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Save document
            doc.save(output_file)
            return True

        except Exception as e:
            print(f"Error creating DOCX file: {str(e)}")
            return False

    def _write_tex(self, content: str, output_file: Path) -> bool:
        """
        Write content to LaTeX file.

        Args:
            content: Text content
            output_file: Path to output file

        Returns:
            True if successful, False otherwise
        """
        try:
            # Create a basic LaTeX document structure
            latex_content = """\\documentclass{article}
\\usepackage[utf8]{inputenc}
\\usepackage{geometry}
\\geometry{a4paper, margin=1in}

\\begin{document}

"""
            # Escape special LaTeX characters
            escaped_content = content.replace('\\', '\\textbackslash{}')
            escaped_content = escaped_content.replace('&', '\\&')
            escaped_content = escaped_content.replace('%', '\\%')
            escaped_content = escaped_content.replace('$', '\\$')
            escaped_content = escaped_content.replace('#', '\\#')
            escaped_content = escaped_content.replace('_', '\\_')
            escaped_content = escaped_content.replace('{', '\\{')
            escaped_content = escaped_content.replace('}', '\\}')
            escaped_content = escaped_content.replace('~', '\\textasciitilde{}')
            escaped_content = escaped_content.replace('^', '\\textasciicircum{}')

            latex_content += escaped_content
            latex_content += "\n\n\\end{document}"

            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(latex_content)

            return True

        except Exception as e:
            print(f"Error creating TEX file: {str(e)}")
            return False

    def _write_txt(self, content: str, output_file: Path) -> bool:
        """
        Write content to plain text file.

        Args:
            content: Text content
            output_file: Path to output file

        Returns:
            True if successful, False otherwise
        """
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(content)
            return True

        except Exception as e:
            print(f"Error creating TXT file: {str(e)}")
            return False

    def _write_md(self, content: str, output_file: Path) -> bool:
        """
        Write content to Markdown file.

        Args:
            content: Text content
            output_file: Path to output file

        Returns:
            True if successful, False otherwise
        """
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(content)
            return True

        except Exception as e:
            print(f"Error creating MD file: {str(e)}")
            return False
